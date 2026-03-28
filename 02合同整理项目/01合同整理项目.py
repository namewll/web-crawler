import os
import docx
import re
import openpyxl

def shx(pattern):
    for i in paragraphs:
        if re.findall(f"{pattern}", i.text):
            value = re.findall(f"{pattern}", i.text)[0]
            content.append(value)

path="./安平市建材制造有限公司销售合同"
new_paths=[]
for i in os.listdir(path):
    new_path=os.path.join(path,i)
    new_paths.append(new_path)

index=0
wb = openpyxl.load_workbook('.\销售合同台账模板.xlsx')
ws = wb.active
for j in new_paths:
    content = []
    # 合同编号
    htnum=re.findall("编号：(.*?)）",j)[0]

    docx_obj=docx.Document(j)
    # 品名
    pinming=docx_obj.tables[0].rows[1].cells[1].text
    content.append(htnum)
    content.append(pinming)
    paragraphs=docx_obj.paragraphs
    shx('乙方：(.*)\t')
    shx('签订日期： (.*)')
    shx('签订地点： (.*)')
    shx(' 授权代表（签名）：(.*)')
    shx(' 联系电话：(.*)')
    shx(' 开户银行：(.*)')
    shx(' 开户账号：(.*)')

    index+=1
    ws.append([index]+content)
wb.save("./销售合同结果.xlsx")
wb.close()





